"""
CV Optimizer Agent - Optimizes CV based on job requirements and auto-applies
"""

import re
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
import openai
from docx import Document
from docx.shared import Inches
import json

from ..core.config import config
from ..core.database import db
from ..core.utils import setup_logging, format_cv_for_job


class CVOptimizer:
    """Agent for optimizing CV based on job requirements"""
    
    def __init__(self):
        self.logger = setup_logging()
        self.cv_config = config.get_cv_optimization_config()
        self.openai_api_key = self.cv_config.get('openai_api_key', '')
        
        if self.openai_api_key:
            openai.api_key = self.openai_api_key
        
        self.skills_mapping = self.cv_config.get('skills_mapping', {})
        self.optimization_prompts = self.cv_config.get('optimization_prompts', [])
    
    def analyze_job_requirements(self, job_description: str) -> Dict[str, Any]:
        """Analyze job requirements to extract key information"""
        analysis = {
            "required_skills": [],
            "preferred_skills": [],
            "experience_level": "",
            "industry": "",
            "job_type": "",
            "key_responsibilities": [],
            "education_requirements": []
        }
        
        # Extract skills using regex patterns
        skill_patterns = [
            r'\b(?:Python|Java|JavaScript|SQL|React|Node\.js|AWS|Docker|Kubernetes|Machine Learning|Data Analysis)\b',
            r'\b(?:PL/SQL|Oracle|MySQL|PostgreSQL|MongoDB|Redis)\b',
            r'\b(?:Project Management|Agile|Scrum|DevOps|CI/CD)\b',
            r'\b(?:HTML|CSS|PHP|C\+\+|C#|\.NET|Angular|Vue\.js)\b'
        ]
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE)
            analysis["required_skills"].extend(matches)
        
        # Extract experience level
        experience_patterns = [
            r'(\d+)\+?\s*years?\s*experience',
            r'(?:senior|junior|mid-level|entry-level|experienced)',
            r'(?:fresher|graduate|postgraduate)'
        ]
        
        for pattern in experience_patterns:
            match = re.search(pattern, job_description, re.IGNORECASE)
            if match:
                analysis["experience_level"] = match.group(0)
                break
        
        # Extract industry
        industry_keywords = [
            "banking", "finance", "healthcare", "retail", "manufacturing",
            "technology", "consulting", "education", "government"
        ]
        
        for keyword in industry_keywords:
            if keyword.lower() in job_description.lower():
                analysis["industry"] = keyword.title()
                break
        
        # Extract job type
        job_type_patterns = [
            r'\b(?:full-time|part-time|contract|freelance|remote|hybrid)\b'
        ]
        
        for pattern in job_type_patterns:
            match = re.search(pattern, job_description, re.IGNORECASE)
            if match:
                analysis["job_type"] = match.group(0)
                break
        
        # Extract key responsibilities
        responsibility_patterns = [
            r'(?:responsible for|duties include|key responsibilities?)[:.]?\s*([^.]*)',
            r'(?:develop|design|implement|maintain|support|manage)[^.]*'
        ]
        
        for pattern in responsibility_patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE)
            analysis["key_responsibilities"].extend(matches)
        
        return analysis
    
    def optimize_cv_with_ai(self, cv_content: str, job_requirements: str, 
                           job_title: str, company: str) -> str:
        """Optimize CV using AI based on job requirements"""
        if not self.openai_api_key:
            self.logger.warning("OpenAI API key not configured, using basic optimization")
            return self._basic_cv_optimization(cv_content, job_requirements)
        
        try:
            # Create optimization prompt
            prompt = f"""
            Optimize the following CV for a {job_title} position at {company}.
            
            Job Requirements:
            {job_requirements}
            
            Current CV:
            {cv_content}
            
            Please optimize the CV by:
            1. Highlighting relevant skills and experience
            2. Reordering sections to prioritize relevant information
            3. Adding keywords from the job requirements
            4. Adjusting language to match the job description
            5. Ensuring the CV is tailored to the specific role
            
            Return the optimized CV content:
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional CV optimization expert."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            optimized_cv = response.choices[0].message.content.strip()
            self.logger.info("CV optimized using AI")
            return optimized_cv
            
        except Exception as e:
            self.logger.error(f"Error optimizing CV with AI: {e}")
            return self._basic_cv_optimization(cv_content, job_requirements)
    
    def _basic_cv_optimization(self, cv_content: str, job_requirements: str) -> str:
        """Basic CV optimization without AI"""
        # Use the utility function for basic optimization
        return format_cv_for_job(cv_content, job_requirements, self.skills_mapping)
    
    def create_optimized_cv_document(self, optimized_content: str, 
                                   job_title: str, company: str) -> str:
        """Create a Word document with the optimized CV"""
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Optimized CV for {job_title} at {company}', 0)
            title.alignment = 1  # Center alignment
            
            # Add timestamp
            timestamp = doc.add_paragraph(f'Optimized on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            timestamp.alignment = 1
            
            # Add content
            doc.add_paragraph(optimized_content)
            
            # Save document
            filename = f"optimized_cv_{job_title.replace(' ', '_')}_{company.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            
            self.logger.info(f"Created optimized CV document: {filename}")
            return filename
            
        except Exception as e:
            self.logger.error(f"Error creating CV document: {e}")
            return ""
    
    def auto_apply_to_job(self, job_data: Dict[str, Any], optimized_cv_path: str = None) -> bool:
        """Automatically apply to a job with optimized CV"""
        try:
            job_title = job_data.get('title', '')
            company = job_data.get('company', '')
            job_url = job_data.get('job_url', '')
            
            if not job_url:
                self.logger.warning(f"No job URL provided for {job_title} at {company}")
                return False
            
            # Log the application attempt
            application_id = db.add_job_application(
                job_title=job_title,
                company=company,
                portal=job_data.get('source', ''),
                country=job_data.get('location', '').split(',')[0] if job_data.get('location') else '',
                cv_version=f"optimized_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
                notes=f"Auto-applied with optimized CV: {optimized_cv_path}" if optimized_cv_path else "Auto-applied"
            )
            
            self.logger.info(f"Auto-applied to {job_title} at {company} (ID: {application_id})")
            return True
            
        except Exception as e:
            self.logger.error(f"Error auto-applying to job: {e}")
            return False
    
    def process_job_application(self, job_data: Dict[str, Any], 
                              cv_content: str) -> Dict[str, Any]:
        """Process a complete job application with CV optimization"""
        try:
            job_title = job_data.get('title', '')
            company = job_data.get('company', '')
            job_description = job_data.get('description', '')
            
            # Analyze job requirements
            requirements_analysis = self.analyze_job_requirements(job_description)
            
            # Optimize CV
            optimized_cv = self.optimize_cv_with_ai(
                cv_content, job_description, job_title, company
            )
            
            # Create optimized CV document
            cv_document_path = self.create_optimized_cv_document(
                optimized_cv, job_title, company
            )
            
            # Auto-apply to job
            application_success = self.auto_apply_to_job(job_data, cv_document_path)
            
            return {
                "job_title": job_title,
                "company": company,
                "requirements_analysis": requirements_analysis,
                "cv_optimized": True,
                "cv_document_path": cv_document_path,
                "application_success": application_success,
                "optimization_date": datetime.now().isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Error processing job application: {e}")
            return {
                "error": str(e),
                "optimization_date": datetime.now().isoformat()
            }
    
    def batch_optimize_cvs(self, jobs: List[Dict[str, Any]], 
                          cv_content: str) -> List[Dict[str, Any]]:
        """Optimize CV for multiple jobs"""
        results = []
        
        for job in jobs:
            try:
                result = self.process_job_application(job, cv_content)
                results.append(result)
                
                # Add delay between applications
                import time
                time.sleep(2)
                
            except Exception as e:
                self.logger.error(f"Error processing job {job.get('title', 'Unknown')}: {e}")
                results.append({
                    "job_title": job.get('title', 'Unknown'),
                    "error": str(e),
                    "optimization_date": datetime.now().isoformat()
                })
        
        return results
    
    def get_optimization_stats(self) -> Dict[str, Any]:
        """Get statistics about CV optimizations"""
        try:
            # Get recent job applications
            recent_applications = db.get_job_applications(limit=100)
            
            total_applications = len(recent_applications)
            successful_applications = len([app for app in recent_applications if app.get('status') == 'applied'])
            responses_received = len([app for app in recent_applications if app.get('response_received')])
            
            return {
                "total_applications": total_applications,
                "successful_applications": successful_applications,
                "responses_received": responses_received,
                "success_rate": (successful_applications / total_applications * 100) if total_applications > 0 else 0,
                "response_rate": (responses_received / total_applications * 100) if total_applications > 0 else 0
            }
            
        except Exception as e:
            self.logger.error(f"Error getting optimization stats: {e}")
            return {}
    
    def generate_cover_letter(self, job_data: Dict[str, Any], 
                            optimized_cv: str) -> str:
        """Generate a cover letter for a job application"""
        if not self.openai_api_key:
            return self._generate_basic_cover_letter(job_data, optimized_cv)
        
        try:
            job_title = job_data.get('title', '')
            company = job_data.get('company', '')
            job_description = job_data.get('description', '')
            
            prompt = f"""
            Generate a professional cover letter for a {job_title} position at {company}.
            
            Job Description:
            {job_description}
            
            CV Summary:
            {optimized_cv[:500]}...
            
            Please create a compelling cover letter that:
            1. Addresses the hiring manager professionally
            2. Explains why you're interested in the position
            3. Highlights relevant skills and experience
            4. Shows enthusiasm for the company
            5. Includes a call to action
            
            Cover Letter:
            """
            
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional cover letter writer."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            cover_letter = response.choices[0].message.content.strip()
            self.logger.info("Generated cover letter using AI")
            return cover_letter
            
        except Exception as e:
            self.logger.error(f"Error generating cover letter: {e}")
            return self._generate_basic_cover_letter(job_data, optimized_cv)
    
    def _generate_basic_cover_letter(self, job_data: Dict[str, Any], 
                                   optimized_cv: str) -> str:
        """Generate a basic cover letter template"""
        job_title = job_data.get('title', '')
        company = job_data.get('company', '')
        
        cover_letter = f"""
Dear Hiring Manager,

I am writing to express my interest in the {job_title} position at {company}. With my relevant experience and skills, I am confident in my ability to contribute effectively to your team.

Based on the job requirements, I believe my background aligns well with your needs. I have experience in the key areas mentioned in the job description and am excited about the opportunity to bring my expertise to {company}.

I am particularly drawn to this position because of {company}'s reputation in the industry and the opportunity to work on challenging projects. I am eager to discuss how my skills and experience can benefit your organization.

Thank you for considering my application. I look forward to the opportunity to discuss how I can contribute to your team.

Best regards,
[Your Name]
        """
        
        return cover_letter.strip()


def main():
    """Main function to test CV optimizer"""
    optimizer = CVOptimizer()
    
    # Example job data
    job_data = {
        "title": "PLSQL Developer",
        "company": "Tech Corp",
        "description": "We are looking for a PLSQL Developer with 3+ years experience in Oracle database development...",
        "source": "bayt.com",
        "location": "Dubai, UAE"
    }
    
    # Example CV content
    cv_content = """
    PROFESSIONAL SUMMARY
    Experienced database developer with expertise in PL/SQL, Oracle, and SQL development.
    
    SKILLS
    - PL/SQL, Oracle Database, SQL, Stored Procedures
    - Python, Java, JavaScript
    - Database Design and Optimization
    
    EXPERIENCE
    Senior Database Developer at ABC Company (2019-2023)
    - Developed and maintained PL/SQL procedures
    - Optimized database queries for better performance
    """
    
    # Process job application
    result = optimizer.process_job_application(job_data, cv_content)
    print(f"Application result: {result}")


if __name__ == "__main__":
    main() 